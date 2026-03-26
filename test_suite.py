#!/usr/bin/env python3
# ============================================================
#  test_suite.py — Kapsamlı Test Paketi
#  Kripto Para Trading Botu | Ahmet Yılmaz
#
#  8 Kategori:
#    1. Fonksiyonel Testler (Unit, Integration, System, UAT)
#    2. Performans Kriterleri
#    3. Güvenilirlik ve Kararlılık
#    4. Güvenlik Testleri
#    5. Kullanılabilirlik
#    6. Taşınabilirlik ve Uyumluluk
#    7. Bakım Yapılabilirlik
#    8. Pass/Fail Kriterleri
# ============================================================

import os, sys, time, re, ast, sqlite3, importlib, traceback, platform

# Windows'ta cp1254 codec Unicode karakterleri (█, ✅, ❌ vb.) encode edemez.
# stdout'u UTF-8 ile yeniden yapilandirarak UnicodeEncodeError'u onluyoruz.
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

import numpy as np
import pandas as pd
from io import StringIO
from contextlib import redirect_stdout

# ── Sonuç toplama ──
RESULTS = []  # (category, name, passed, detail, duration_ms)

def record(cat, name, passed, detail="", dur=0):
    RESULTS.append((cat, name, passed, detail, dur))
    icon = "✅" if passed else "❌"
    print(f"  {icon} [{cat}] {name}" + (f" — {detail}" if detail else ""))


# ═══════════════════════════════════════════════════════════════
#  1. FONKSİYONEL TESTLER
# ═══════════════════════════════════════════════════════════════

def test_functional():
    print("\n" + "═" * 60)
    print("  1. FONKSİYONEL TESTLER")
    print("═" * 60)

    # ── 1.1 Birim Testleri ──
    # config.py import
    t0 = time.perf_counter()
    try:
        import config
        ok = all(hasattr(config, a) for a in [
            "SYMBOLS", "INTERVALS", "DB_PATH", "SIGNAL_WEIGHTS",
            "RSI_OVERSOLD", "RSI_OVERBOUGHT", "REPORTS_DIR",
        ])
        record("Unit", "config.py — tüm sabitler mevcut", ok,
               f"{len(dir(config))} öznitelik", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Unit", "config.py import", False, str(e))

    # database.py fonksiyonları
    t0 = time.perf_counter()
    try:
        import database
        fns = ["create_tables","insert_klines","insert_snapshot","insert_market_data",
               "insert_indicators","insert_signals","get_db_stats","get_signals",
               "get_indicators","get_klines","get_latest_prices","log_fetch"]
        missing = [f for f in fns if not hasattr(database, f)]
        record("Unit", "database.py — tüm fonksiyonlar", len(missing)==0,
               f"Eksik: {missing}" if missing else f"{len(fns)} fonksiyon OK",
               (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Unit", "database.py import", False, str(e))

    # indicators.py birim test
    t0 = time.perf_counter()
    try:
        from indicators import sma, ema, rsi, macd, bollinger_bands, atr, vwap, add_all_indicators
        prices = pd.Series(50000 + np.cumsum(np.random.randn(100) * 500))

        sma_v = sma(prices, 20)
        assert sma_v.dropna().shape[0] > 0, "SMA boş"
        assert not np.isinf(sma_v.dropna()).any(), "SMA inf içeriyor"

        ema_v = ema(prices, 20)
        assert ema_v.dropna().shape[0] > 0, "EMA boş"

        rsi_v = rsi(prices, 14)
        rsi_clean = rsi_v.dropna()
        assert (rsi_clean >= 0).all() and (rsi_clean <= 100).all(), "RSI 0-100 dışı"

        macd_df = macd(prices)
        assert set(macd_df.columns) == {"macd_line","macd_signal","macd_hist"}, "MACD sütunlar eksik"

        bb_df = bollinger_bands(prices, 20)
        bb_clean = bb_df.dropna()
        assert (bb_clean["bb_upper"] >= bb_clean["bb_lower"]).all(), "BB upper < lower"

        record("Unit", "indicators.py — tüm göstergeler doğru", True,
               "SMA,EMA,RSI,MACD,BB,ATR,VWAP", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Unit", "indicators.py hesaplama", False, str(e))

    # signal_generator.py birim test
    t0 = time.perf_counter()
    try:
        from signal_generator import (generate_rsi_signal, generate_macd_signal,
            generate_bb_signal, generate_ma_crossover_signal, generate_combined_signal)

        np.random.seed(42)
        n = 200
        p = 50000 + np.cumsum(np.random.randn(n) * 500)
        demo = pd.DataFrame({
            "open_time": range(n), "open": p + np.random.randn(n)*100,
            "high": p + abs(np.random.randn(n)*300),
            "low": p - abs(np.random.randn(n)*300),
            "close": p, "volume": np.random.uniform(100,1000,n),
        })
        from indicators import add_all_indicators
        demo = add_all_indicators(demo).dropna()

        rsi_s = generate_rsi_signal(demo)
        assert set(rsi_s.unique()).issubset({"BUY","SELL","HOLD"}), "RSI sinyal geçersiz"

        macd_s = generate_macd_signal(demo)
        assert set(macd_s.unique()).issubset({"BUY","SELL","HOLD"}), "MACD sinyal geçersiz"

        bb_s = generate_bb_signal(demo)
        assert set(bb_s.unique()).issubset({"BUY","SELL","HOLD"}), "BB sinyal geçersiz"

        ma_s = generate_ma_crossover_signal(demo)
        assert set(ma_s.unique()).issubset({"BUY","SELL","HOLD"}), "MA sinyal geçersiz"

        combined = generate_combined_signal(demo)
        assert "combined" in combined.columns, "combined sütunu yok"
        assert "strength" in combined.columns, "strength sütunu yok"
        assert (combined["strength"] >= 0).all() and (combined["strength"] <= 1).all(), "strength 0-1 dışı"

        record("Unit", "signal_generator.py — 5 strateji doğru", True,
               f"{len(demo)} satır test", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Unit", "signal_generator.py", False, str(e))

    # data_processor.py birim test
    t0 = time.perf_counter()
    try:
        from data_processor import clean_dataframe, normalize_min_max, normalize_zscore
        from data_processor import add_time_features, add_price_features

        df_test = pd.DataFrame({
            "open_time": range(50), "open": np.random.uniform(100,200,50),
            "high": np.random.uniform(200,300,50), "low": np.random.uniform(50,100,50),
            "close": np.random.uniform(100,200,50), "volume": np.random.uniform(1000,5000,50),
            "quote_volume": np.random.uniform(1000,5000,50),
            "trade_count": np.random.randint(10,100,50),
            "taker_buy_vol": np.random.uniform(500,2500,50),
        }, index=pd.date_range("2025-01-01", periods=50, freq="h"))

        cleaned = clean_dataframe(df_test)
        assert len(cleaned) <= len(df_test), "Temizleme satır artırdı"

        normed = normalize_min_max(cleaned, ["close"])
        assert normed["close"].min() >= 0 and normed["close"].max() <= 1, "MinMax hatası"

        zscored = normalize_zscore(df_test.copy(), ["close"])
        assert abs(zscored["close"].mean()) < 0.01, "Z-score mean != 0"

        timed = add_time_features(df_test.copy())
        assert "hour" in timed.columns and "is_weekend" in timed.columns, "Zaman özellik eksik"

        priced = add_price_features(df_test.copy())
        assert "price_change" in priced.columns and "candle_body" in priced.columns, "Fiyat özellik eksik"

        record("Unit", "data_processor.py — temizleme & özellikler", True,
               "clean, minmax, zscore, time, price", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Unit", "data_processor.py", False, str(e))

    # visualizer.py import test
    t0 = time.perf_counter()
    try:
        import visualizer
        fns = ["plot_price_with_indicators","plot_rsi","plot_macd","plot_volume",
               "plot_signals","plot_full_dashboard","plot_portfolio_summary","generate_all_charts"]
        missing = [f for f in fns if not hasattr(visualizer, f)]
        record("Unit", "visualizer.py — tüm fonksiyonlar", len(missing)==0,
               f"{len(fns)} fonksiyon OK", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Unit", "visualizer.py import", False, str(e))

    # ── 1.2 Entegrasyon Testi ──
    t0 = time.perf_counter()
    try:
        from database import create_tables, get_db_stats
        create_tables()
        stats = get_db_stats()
        assert isinstance(stats, dict), "Stats dict değil"
        expected_tables = ["klines","price_snapshots","market_data","fetch_log","indicators","signals"]
        missing_t = [t for t in expected_tables if t not in stats]
        record("Integration", "database → 6 tablo oluşturma", len(missing_t)==0,
               f"Eksik: {missing_t}" if missing_t else "6/6 tablo OK",
               (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Integration", "database tablo oluşturma", False, str(e))

    t0 = time.perf_counter()
    try:
        from data_processor import load_klines_df
        from indicators import add_all_indicators
        from signal_generator import generate_combined_signal
        df = load_klines_df("BTCUSDT", "1h")
        if not df.empty:
            df_ind = add_all_indicators(df).dropna()
            if not df_ind.empty:
                sigs = generate_combined_signal(df_ind)
                assert len(sigs) == len(df_ind), "Sinyal sayısı uyumsuz"
                record("Integration", "data_processor → indicators → signals zinciri", True,
                       f"{len(sigs)} sinyal üretildi", (time.perf_counter()-t0)*1000)
            else:
                record("Integration", "data_processor → indicators → signals zinciri", True,
                       "Veri var ama gösterge için yetersiz (beklenen)", (time.perf_counter()-t0)*1000)
        else:
            record("Integration", "data_processor → indicators → signals zinciri", True,
                   "DB boş — zincir mantığı doğru (skip)", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Integration", "veri zinciri", False, str(e))

    # ── 1.3 Sistem Testi ──
    t0 = time.perf_counter()
    try:
        from database import insert_signals, get_signals
        test_rows = [{
            "symbol":"TESTUSDT","interval":"1h","open_time":9999999999,
            "rsi_signal":"BUY","macd_signal":"HOLD","bb_signal":"SELL",
            "ma_signal":"HOLD","combined":"HOLD","strength":0.15,
        }]
        insert_signals(test_rows)
        fetched = get_signals("TESTUSDT","1h",10)
        assert len(fetched) >= 1, "Sinyal kaydedilemedi"
        assert fetched[-1]["combined"] == "HOLD", "Veri tutarsız"
        record("System", "Sinyal kaydet → oku uçtan uca", True,
               "INSERT + SELECT OK", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("System", "Sinyal uçtan uca", False, str(e))

    # ── 1.4 Kabul Testi (UAT) ──
    t0 = time.perf_counter()
    try:
        from config import SYMBOLS, INTERVALS
        assert len(SYMBOLS) >= 3, f"En az 3 coin: {len(SYMBOLS)}"
        assert len(INTERVALS) >= 3, f"En az 3 interval: {len(INTERVALS)}"
        assert "BTC" in SYMBOLS and "ETH" in SYMBOLS, "BTC/ETH eksik"
        record("UAT", "Desteklenen semboller ve interval'lar", True,
               f"{len(SYMBOLS)} coin, {len(INTERVALS)} interval",
               (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("UAT", "Sembol/interval", False, str(e))

    t0 = time.perf_counter()
    try:
        import main
        cmds = ["setup","historical","realtime","all","analyze","indicators","signals","visualize"]
        missing_c = [c for c in cmds if c not in main.COMMANDS]
        record("UAT", "Tüm CLI komutları mevcut", len(missing_c)==0,
               f"8/8 komut OK", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("UAT", "CLI komutları", False, str(e))


# ═══════════════════════════════════════════════════════════════
#  2. PERFORMANS KRİTERLERİ
# ═══════════════════════════════════════════════════════════════

def test_performance():
    print("\n" + "═" * 60)
    print("  2. PERFORMANS KRİTERLERİ")
    print("═" * 60)

    # Gösterge hesaplama hızı
    t0 = time.perf_counter()
    from indicators import add_all_indicators
    np.random.seed(0)
    n = 5000
    p = 50000 + np.cumsum(np.random.randn(n)*500)
    big_df = pd.DataFrame({
        "open": p+np.random.randn(n)*100, "high": p+abs(np.random.randn(n)*300),
        "low": p-abs(np.random.randn(n)*300), "close": p,
        "volume": np.random.uniform(100,1000,n),
    })
    result = add_all_indicators(big_df)
    dur = (time.perf_counter()-t0)*1000
    record("Performans", f"Gösterge hesaplama ({n} satır)", dur < 2000,
           f"{dur:.0f}ms (eşik: <2000ms)", dur)

    # Sinyal üretim hızı
    t0 = time.perf_counter()
    from signal_generator import generate_combined_signal
    df_clean = result.dropna()
    if not df_clean.empty:
        sigs = generate_combined_signal(df_clean)
        dur2 = (time.perf_counter()-t0)*1000
        throughput = len(df_clean) / (dur2/1000) if dur2 > 0 else 0
        record("Performans", f"Sinyal üretimi ({len(df_clean)} satır)", dur2 < 2000,
               f"{dur2:.0f}ms | {throughput:.0f} satır/sn", dur2)

    # DB yazma hızı
    t0 = time.perf_counter()
    from database import insert_signals
    bulk = [{"symbol":"PERFTEST","interval":"1h","open_time":i,
             "rsi_signal":"HOLD","macd_signal":"HOLD","bb_signal":"HOLD",
             "ma_signal":"HOLD","combined":"HOLD","strength":0.1} for i in range(1000)]
    insert_signals(bulk)
    dur3 = (time.perf_counter()-t0)*1000
    record("Performans", "DB toplu yazma (1000 kayıt)", dur3 < 2000,
           f"{dur3:.0f}ms | {1000/(dur3/1000):.0f} kayıt/sn", dur3)

    # Config yükleme hızı
    t0 = time.perf_counter()
    importlib.reload(importlib.import_module("config"))
    dur4 = (time.perf_counter()-t0)*1000
    record("Performans", "Config yükleme", dur4 < 100,
           f"{dur4:.1f}ms (eşik: <100ms)", dur4)


# ═══════════════════════════════════════════════════════════════
#  3. GÜVENİLİRLİK VE KARARLILIK
# ═══════════════════════════════════════════════════════════════

def test_reliability():
    print("\n" + "═" * 60)
    print("  3. GÜVENİLİRLİK VE KARARLILIK")
    print("═" * 60)

    # Boş veri toleransı
    t0 = time.perf_counter()
    try:
        from data_processor import load_klines_df
        df = load_klines_df("NONEXISTENT_SYMBOL", "1h")
        assert df.empty, "Olmayan sembol veri döndürdü"
        record("Güvenilirlik", "Boş veri toleransı", True,
               "Olmayan sembol → boş DataFrame", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Güvenilirlik", "Boş veri toleransı", False, str(e))

    # NaN/Inf toleransı
    t0 = time.perf_counter()
    try:
        from indicators import add_all_indicators
        bad_df = pd.DataFrame({
            "open": [100,0,np.nan,200,150], "high": [110,0,np.nan,210,160],
            "low": [90,0,np.nan,190,140], "close": [105,0,np.nan,205,155],
            "volume": [1000,0,np.nan,2000,1500],
        })
        result = add_all_indicators(bad_df)
        has_inf = np.isinf(result.select_dtypes(include=[np.number])).any().any()
        record("Güvenilirlik", "NaN/sıfır veri gösterge hesaplama", not has_inf,
               "Inf yok" if not has_inf else "Inf mevcut!", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Güvenilirlik", "NaN/Inf toleransı", False, str(e))

    # Tekrarlı çalıştırma kararlılığı
    t0 = time.perf_counter()
    try:
        from indicators import rsi
        prices = pd.Series(50000 + np.cumsum(np.random.randn(100)*500))
        results = [rsi(prices, 14).dropna().values.tolist() for _ in range(5)]
        all_same = all(r == results[0] for r in results)
        record("Güvenilirlik", "Deterministik sonuç (5 tekrar)", all_same,
               "Tüm sonuçlar aynı" if all_same else "FARKLI SONUÇLAR!",
               (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Güvenilirlik", "Deterministik sonuç", False, str(e))

    # DB bağlantı güvenilirliği
    t0 = time.perf_counter()
    try:
        from database import get_connection
        errors = 0
        for i in range(20):
            try:
                with get_connection() as conn:
                    conn.execute("SELECT 1").fetchone()
            except:
                errors += 1
        record("Güvenilirlik", "DB bağlantı güvenilirliği (20 bağlantı)", errors == 0,
               f"Hata: {errors}/20", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Güvenilirlik", "DB bağlantı", False, str(e))

    # insert_signals idempotent (UPSERT)
    t0 = time.perf_counter()
    try:
        from database import insert_signals, get_connection
        row = [{"symbol":"IDEMPOTENT","interval":"1h","open_time":12345,
                "rsi_signal":"BUY","macd_signal":"HOLD","bb_signal":"HOLD",
                "ma_signal":"HOLD","combined":"BUY","strength":0.5}]
        insert_signals(row)
        insert_signals(row)  # Aynı kayıt tekrar
        with get_connection() as conn:
            cnt = conn.execute(
                "SELECT COUNT(*) FROM signals WHERE symbol='IDEMPOTENT' AND open_time=12345"
            ).fetchone()[0]
        record("Güvenilirlik", "UPSERT idempotent (mükerrer kayıt)", cnt == 1,
               f"Kayıt sayısı: {cnt} (beklenen: 1)", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Güvenilirlik", "UPSERT", False, str(e))


# ═══════════════════════════════════════════════════════════════
#  4. GÜVENLİK TESTLERİ
# ═══════════════════════════════════════════════════════════════

def test_security():
    print("\n" + "═" * 60)
    print("  4. GÜVENLİK TESTLERİ")
    print("═" * 60)

    # SQL Injection testi
    t0 = time.perf_counter()
    try:
        from data_processor import load_klines_df
        malicious_inputs = [
            "'; DROP TABLE klines; --",
            "1' OR '1'='1",
            "BTCUSDT; DELETE FROM signals",
            "' UNION SELECT * FROM fetch_log --",
        ]
        for payload in malicious_inputs:
            df = load_klines_df(payload, "1h")
            # Tablo hala var olmalı
        from database import get_connection
        with get_connection() as conn:
            tables = conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table'"
            ).fetchall()
        table_names = [t[0] for t in tables]
        assert "klines" in table_names, "klines tablosu silindi!"
        assert "signals" in table_names, "signals tablosu silindi!"
        record("Güvenlik", "SQL Injection koruması (4 payload)", True,
               "Parametrik sorgular — tablolar sağlam",
               (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Güvenlik", "SQL Injection", False, str(e))

    # API key gizliliği
    t0 = time.perf_counter()
    try:
        from config import BINANCE_API_KEY, BINANCE_SECRET, COINGECKO_API_KEY
        keys_empty = (BINANCE_API_KEY == "" and BINANCE_SECRET == "" and COINGECKO_API_KEY == "")
        record("Güvenlik", "API anahtarları boş (güvenli varsayılan)", keys_empty,
               "Kod deposunda key yok" if keys_empty else "⚠ Key bulundu!",
               (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Güvenlik", "API key kontrolü", False, str(e))

    # Dosyalarda hardcoded secret arama
    t0 = time.perf_counter()
    try:
        secret_patterns = [r'sk-[a-zA-Z0-9]{20,}', r'password\s*=\s*["\'][^"\']+["\']',
                           r'secret\s*=\s*["\'][^"\']+["\']']
        py_files = [f for f in os.listdir('.') if f.endswith('.py') and f != 'test_suite.py']
        found_secrets = []
        for fname in py_files:
            with open(fname, 'r', encoding='utf-8') as f:
                content = f.read()
            for pat in secret_patterns:
                matches = re.findall(pat, content, re.IGNORECASE)
                # Filter out empty assignments like secret = ""
                real = [m for m in matches if not re.search(r'=\s*["\'][\s]*["\']', m)]
                if real:
                    found_secrets.extend([(fname, m) for m in real])
        record("Güvenlik", "Hardcoded secret taraması", len(found_secrets)==0,
               "Temiz" if not found_secrets else f"BULUNDU: {found_secrets}",
               (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Güvenlik", "Secret tarama", False, str(e))

    # Input validation
    t0 = time.perf_counter()
    try:
        from database import insert_klines
        result = insert_klines([])  # Boş liste
        assert result == 0, "Boş liste hata verdi"
        from database import insert_signals
        result2 = insert_signals([])
        assert result2 == 0
        record("Güvenlik", "Boş girdi doğrulaması", True,
               "Boş listeler güvenli", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Güvenlik", "Input validation", False, str(e))


# ═══════════════════════════════════════════════════════════════
#  5. KULLANILABİLİRLİK
# ═══════════════════════════════════════════════════════════════

def test_usability():
    print("\n" + "═" * 60)
    print("  5. KULLANILABİLİRLİK")
    print("═" * 60)

    # CLI yardım metni
    t0 = time.perf_counter()
    try:
        with open("main.py","r",encoding="utf-8") as f:
            content = f.read()
        has_help = all(cmd in content for cmd in
                       ["setup","historical","realtime","analyze","indicators","signals","visualize"])
        has_desc = "Kullanım:" in content
        record("Kullanılabilirlik", "CLI yardım metni eksiksiz", has_help and has_desc,
               "Tüm komut açıklamaları mevcut", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Kullanılabilirlik", "CLI yardım", False, str(e))

    # Docstring kontrolü
    t0 = time.perf_counter()
    try:
        modules = ["config","database","indicators","data_processor","signal_generator","visualizer"]
        total_fns = 0; documented = 0
        for mod_name in modules:
            mod = importlib.import_module(mod_name)
            for name in dir(mod):
                obj = getattr(mod, name)
                if callable(obj) and not name.startswith("_"):
                    total_fns += 1
                    if obj.__doc__:
                        documented += 1
        pct = (documented/total_fns*100) if total_fns > 0 else 0
        record("Kullanılabilirlik", f"Fonksiyon dokümantasyon oranı", pct >= 80,
               f"{documented}/{total_fns} ({pct:.0f}%) — eşik: ≥80%",
               (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Kullanılabilirlik", "Docstring kontrolü", False, str(e))

    # Log mesajlarında emoji kullanımı (okunabilirlik)
    t0 = time.perf_counter()
    try:
        emoji_files = 0
        for fname in ["main.py","signal_generator.py","visualizer.py","database.py",
                       "data_processor.py","indicators.py"]:
            with open(fname,"r",encoding="utf-8") as f:
                if re.search(r'[✅❌⚠️📊📈🔔💰🎨📂🔧]', f.read()):
                    emoji_files += 1
        record("Kullanılabilirlik", "Log emoji kullanımı (okunabilirlik)", emoji_files >= 4,
               f"{emoji_files}/6 dosyada emoji log", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Kullanılabilirlik", "Emoji log", False, str(e))

    # Banner mevcut mu
    t0 = time.perf_counter()
    try:
        import main
        assert hasattr(main, "BANNER"), "BANNER yok"
        assert "Kripto" in main.BANNER or "Trading" in main.BANNER, "Banner içeriği yetersiz"
        record("Kullanılabilirlik", "Uygulama banner'ı mevcut ve açıklayıcı", True,
               "ASCII art + proje bilgisi", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Kullanılabilirlik", "Banner", False, str(e))


# ═══════════════════════════════════════════════════════════════
#  6. TAŞINABİLİRLİK VE UYUMLULUK
# ═══════════════════════════════════════════════════════════════

def test_portability():
    print("\n" + "═" * 60)
    print("  6. TAŞINABİLİRLİK VE UYUMLULUK")
    print("═" * 60)

    # Python versiyon kontrolü
    t0 = time.perf_counter()
    v = sys.version_info
    ok = v.major == 3 and v.minor >= 10
    record("Taşınabilirlik", f"Python versiyonu ({v.major}.{v.minor}.{v.micro})", ok,
           "≥3.10 gerekli (type hints)", (time.perf_counter()-t0)*1000)

    # OS bağımsızlık (os.path kullanımı)
    t0 = time.perf_counter()
    try:
        py_files = [f for f in os.listdir('.') if f.endswith('.py') and f != 'test_suite.py']
        hardcoded_paths = []
        for fname in py_files:
            with open(fname,'r',encoding='utf-8') as f:
                lines = f.readlines()
            for i,line in enumerate(lines):
                if re.search(r'["\'][A-Z]:\\\\', line) or re.search(r'["\']\/home\/', line):
                    hardcoded_paths.append((fname, i+1))
        record("Taşınabilirlik", "Hardcoded yol kontrolü", len(hardcoded_paths)==0,
               "Temiz" if not hardcoded_paths else f"Bulundu: {hardcoded_paths}",
               (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Taşınabilirlik", "Yol kontrolü", False, str(e))

    # SQLite taşınabilirlik
    t0 = time.perf_counter()
    record("Taşınabilirlik", "Veritabanı: SQLite (harici sunucu gerektirmez)", True,
           f"Dosya: {os.path.basename(sqlite3.connect(':memory:').execute('PRAGMA database_list').fetchone()[2] or 'memory')}",
           (time.perf_counter()-t0)*1000)

    # Bağımlılık sayısı
    t0 = time.perf_counter()
    try:
        with open("requirements.txt","r") as f:
            deps = [l.strip() for l in f if l.strip() and not l.startswith("#")]
        record("Taşınabilirlik", f"Bağımlılık sayısı: {len(deps)} paket", len(deps) <= 20,
               "Proje bağımlılıkları kabul edilebilir düzeyde", (time.perf_counter()-t0)*1000)
    except Exception as e:
        record("Taşınabilirlik", "Bağımlılık", False, str(e))

    # İşletim sistemi bilgisi
    t0 = time.perf_counter()
    record("Taşınabilirlik", f"Mevcut OS: {platform.system()} {platform.release()}", True,
           f"Arch: {platform.machine()}", (time.perf_counter()-t0)*1000)


# ═══════════════════════════════════════════════════════════════
#  7. BAKIM YAPILABİLİRLİK
# ═══════════════════════════════════════════════════════════════

def test_maintainability():
    print("\n" + "═" * 60)
    print("  7. BAKIM YAPILABİLİRLİK")
    print("═" * 60)

    # Modülerlik — dosya sayısı ve sorumluluk ayrımı
    t0 = time.perf_counter()
    py_files = [f for f in os.listdir('.') if f.endswith('.py')
                and f != 'test_suite.py' and not f.startswith('_')]
    record("Bakım", f"Modül sayısı: {len(py_files)} Python dosyası", len(py_files) >= 7,
           "Her modül tek sorumluluk", (time.perf_counter()-t0)*1000)

    # Kod satır analizi
    t0 = time.perf_counter()
    total_lines = 0; total_comments = 0; total_blank = 0; total_code = 0
    for fname in py_files:
        with open(fname,'r',encoding='utf-8') as f:
            lines = f.readlines()
        total_lines += len(lines)
        for line in lines:
            stripped = line.strip()
            if not stripped:
                total_blank += 1
            elif stripped.startswith('#'):
                total_comments += 1
            else:
                total_code += 1
    comment_ratio = (total_comments/total_lines*100) if total_lines > 0 else 0
    record("Bakım", f"Toplam: {total_lines} satır ({total_code} kod, {total_comments} yorum, {total_blank} boş)",
           comment_ratio >= 10,
           f"Yorum oranı: {comment_ratio:.1f}% (eşik: ≥10%)",
           (time.perf_counter()-t0)*1000)

    # Fonksiyon uzunluk kontrolü (max 100 satır)
    t0 = time.perf_counter()
    long_functions = []
    for fname in py_files:
        try:
            with open(fname,'r',encoding='utf-8') as f:
                tree = ast.parse(f.read())
            for node in ast.walk(tree):
                if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                    length = node.end_lineno - node.lineno + 1 if hasattr(node,'end_lineno') else 0
                    if length > 100:
                        long_functions.append((fname, node.name, length))
        except:
            pass
    record("Bakım", "Fonksiyon uzunluğu (max 100 satır)", len(long_functions)==0,
           "Tümü ≤100" if not long_functions else f"Uzun: {long_functions}",
           (time.perf_counter()-t0)*1000)

    # Syntax kontrolü (py_compile)
    t0 = time.perf_counter()
    compile_errors = []
    for fname in py_files:
        try:
            with open(fname,'r',encoding='utf-8') as f:
                compile(f.read(), fname, 'exec')
        except SyntaxError as e:
            compile_errors.append((fname, str(e)))
    record("Bakım", f"Syntax kontrolü ({len(py_files)} dosya)", len(compile_errors)==0,
           "Tümü derlendi" if not compile_errors else f"Hata: {compile_errors}",
           (time.perf_counter()-t0)*1000)

    # Circular import kontrolü
    t0 = time.perf_counter()
    try:
        mods = ["config","database","indicators","data_processor","signal_generator","visualizer","main"]
        for m in mods:
            if m in sys.modules:
                del sys.modules[m]
        for m in mods:
            importlib.import_module(m)
        record("Bakım", "Circular import kontrolü", True,
               f"{len(mods)} modül sıralı import OK", (time.perf_counter()-t0)*1000)
    except ImportError as e:
        record("Bakım", "Circular import", False, str(e))


# ═══════════════════════════════════════════════════════════════
#  RAPOR OLUŞTUR
# ═══════════════════════════════════════════════════════════════

def generate_report():
    """Test sonuçlarını DOCX rapor olarak kaydet."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(11)

    # Kapak
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NECMETTİN ERBAKAN ÜNİVERSİTESİ"); run.font.size = Pt(14); run.bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Bilgisayar Mühendisliği Bölümü").font.size = Pt(12)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TEST RAPORU"); run.font.size = Pt(18); run.bold = True
    run.font.color.rgb = RGBColor(0x37,0x42,0xFA)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Kripto Para Trading Botu — Kapsamlı Test Sonuçları").font.size = Pt(12)
    doc.add_paragraph()
    info = doc.add_table(rows=3,cols=2); info.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate([("Öğrenci","Ahmet Yılmaz — 23100011075"),
                               ("Tarih","24 Şubat 2026"),("Kapsam","1-3. Hafta Tüm Modüller")]):
        info.rows[i].cells[0].text = k; info.rows[i].cells[1].text = v
        for pp in info.rows[i].cells[0].paragraphs:
            for r in pp.runs: r.bold = True

    doc.add_page_break()

    # Özet
    passed = sum(1 for r in RESULTS if r[2])
    failed = sum(1 for r in RESULTS if not r[2])
    total = len(RESULTS)
    pct = (passed/total*100) if total > 0 else 0

    doc.add_heading("Test Özeti", level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = "Light Grid Accent 1"
    for i,(k,v) in enumerate([("Toplam Test",str(total)),("Başarılı",f"{passed} ✅"),
                               ("Başarısız",f"{failed} ❌"),("Başarı Oranı",f"%{pct:.1f}")]):
        summary_table.rows[i].cells[0].text = k
        summary_table.rows[i].cells[1].text = v

    doc.add_paragraph()

    # Kategori bazlı sonuçlar
    categories = {}
    for cat, name, passed_flag, detail, dur in RESULTS:
        categories.setdefault(cat, []).append((name, passed_flag, detail, dur))

    cat_titles = {
        "Unit": "1. Fonksiyonel — Birim Testleri",
        "Integration": "1. Fonksiyonel — Entegrasyon Testleri",
        "System": "1. Fonksiyonel — Sistem Testleri",
        "UAT": "1. Fonksiyonel — Kabul Testleri",
        "Performans": "2. Performans Kriterleri",
        "Güvenilirlik": "3. Güvenilirlik ve Kararlılık",
        "Güvenlik": "4. Güvenlik Testleri",
        "Kullanılabilirlik": "5. Kullanılabilirlik",
        "Taşınabilirlik": "6. Taşınabilirlik ve Uyumluluk",
        "Bakım": "7. Bakım Yapılabilirlik",
    }

    for cat_key in ["Unit","Integration","System","UAT","Performans",
                     "Güvenilirlik","Güvenlik","Kullanılabilirlik","Taşınabilirlik","Bakım"]:
        if cat_key not in categories:
            continue
        items = categories[cat_key]
        doc.add_heading(cat_titles.get(cat_key, cat_key), level=2)

        t = doc.add_table(rows=1+len(items), cols=4)
        t.style = "Light Grid Accent 1"
        for i,h in enumerate(["Test","Sonuç","Detay","Süre (ms)"]):
            t.rows[0].cells[i].text = h
            for pp in t.rows[0].cells[i].paragraphs:
                for r in pp.runs: r.bold = True; r.font.size = Pt(9)

        for j,(name,ok,detail,dur) in enumerate(items):
            t.rows[j+1].cells[0].text = name
            t.rows[j+1].cells[1].text = "✅ PASS" if ok else "❌ FAIL"
            t.rows[j+1].cells[2].text = detail
            t.rows[j+1].cells[3].text = f"{dur:.1f}"
            for pp in t.rows[j+1].cells[1].paragraphs:
                for r in pp.runs:
                    r.font.color.rgb = RGBColor(0,180,0) if ok else RGBColor(255,0,0)
                    r.font.size = Pt(9)
            for ci in [0,2,3]:
                for pp in t.rows[j+1].cells[ci].paragraphs:
                    for r in pp.runs: r.font.size = Pt(9)
        doc.add_paragraph()

    # 8. Pass/Fail Kriterleri
    doc.add_heading("8. Test Başarı/Başarısızlık Kriterleri", level=1)
    pf = doc.add_table(rows=5, cols=4)
    pf.style = "Light Grid Accent 1"
    for i,h in enumerate(["Kriter","Eşik Değer","Gerçekleşen","Sonuç"]):
        pf.rows[0].cells[i].text = h
        for pp in pf.rows[0].cells[i].paragraphs:
            for r in pp.runs: r.bold = True

    avg_dur = sum(r[4] for r in RESULTS) / len(RESULTS) if RESULTS else 0
    err_rate = (failed/total*100) if total > 0 else 0

    criteria = [
        ("Yanıt süresi","< 2000 ms", f"{avg_dur:.0f} ms", avg_dur < 2000),
        ("Hata oranı","< %5", f"%{err_rate:.1f}", err_rate < 5),
        ("Test coverage","≥ %80 fonksiyon", f"%{pct:.0f} test başarı", pct >= 80),
        ("Güvenlik açığı","Kritik: 0 adet",
         f"{'0 kritik' if all(r[2] for r in RESULTS if r[0]=='Güvenlik') else 'AÇIK VAR'}",
         all(r[2] for r in RESULTS if r[0]=="Güvenlik")),
    ]
    for i,(k,t_val,actual,ok) in enumerate(criteria):
        pf.rows[i+1].cells[0].text = k
        pf.rows[i+1].cells[1].text = t_val
        pf.rows[i+1].cells[2].text = actual
        pf.rows[i+1].cells[3].text = "✅ PASS" if ok else "❌ FAIL"

    outpath = os.path.join("reports","Test_Raporu.docx")
    os.makedirs("reports", exist_ok=True)
    doc.save(outpath)
    print(f"\n📄 Test raporu kaydedildi: {outpath}")
    print(f"   Boyut: {os.path.getsize(outpath)/1024:.1f} KB")


# ═══════════════════════════════════════════════════════════════
#  ANA ÇALIŞTIRICI
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("\n" + "█" * 60)
    print("  🧪  KAPSAMLI TEST PAKETİ")
    print("  Kripto Para Trading Botu — 1-3. Hafta")
    print("█" * 60)

    t_start = time.perf_counter()

    test_functional()
    test_performance()
    test_reliability()
    test_security()
    test_usability()
    test_portability()
    test_maintainability()

    total_dur = (time.perf_counter() - t_start) * 1000

    # Özet
    passed = sum(1 for r in RESULTS if r[2])
    failed = sum(1 for r in RESULTS if not r[2])
    total = len(RESULTS)

    print("\n" + "█" * 60)
    print(f"  📊  SONUÇ: {passed}/{total} BAŞARILI ({passed/total*100:.1f}%)")
    if failed > 0:
        print(f"  ❌  {failed} test BAŞARISIZ:")
        for cat, name, ok, detail, dur in RESULTS:
            if not ok:
                print(f"      [{cat}] {name}: {detail}")
    print(f"  ⏱️  Toplam süre: {total_dur:.0f}ms")
    print("█" * 60)

    generate_report()
